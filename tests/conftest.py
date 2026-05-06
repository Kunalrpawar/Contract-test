"""
Pytest Fixtures and Configuration
"""

import os
import tempfile
import pytest
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker, Session
from fastapi.testclient import TestClient
from app.database import Base
from app.main import app
from app.database import get_db
from app.config import settings

# Use in-memory SQLite for testing
SQLALCHEMY_DATABASE_URL = "sqlite:///:memory:"

engine = create_engine(
    SQLALCHEMY_DATABASE_URL,
    connect_args={"check_same_thread": False},
)

TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)


@pytest.fixture(scope="function")
def db():
    """Create a fresh database for each test"""
    Base.metadata.create_all(bind=engine)
    yield TestingSessionLocal()
    Base.metadata.drop_all(bind=engine)


@pytest.fixture(scope="function")
def client(db: Session):
    """Create test client with test database"""
    def override_get_db():
        yield db
    
    app.dependency_overrides[get_db] = override_get_db
    yield TestClient(app)
    app.dependency_overrides.clear()


@pytest.fixture(scope="function")
def sample_pdf_file():
    """Create a temporary PDF file for testing"""
    # Create a simple PDF-like file
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        # Minimal PDF structure
        pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> >> >> /MediaBox [0 0 612 792] /Contents 4 0 R >>
endobj
4 0 obj
<< >>
stream
BT
/F1 12 Tf
100 700 Td
(Sample Contract Content) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000273 00000 n 
trailer
<< /Size 5 /Root 1 0 R >>
startxref
353
%%EOF
"""
        f.write(pdf_content)
        f.flush()
        yield f.name
    
    # Cleanup
    if os.path.exists(f.name):
        os.remove(f.name)


@pytest.fixture(scope="function")
def sample_docx_file():
    """Create a temporary DOCX file for testing"""
    from docx import Document
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        doc.add_paragraph("SLA Terms: 99.9% uptime required")
        doc.add_paragraph("Payment Terms: Net 30 days")
        doc.add_paragraph("Termination: 30 days written notice")
        doc.add_paragraph("Confidentiality: 5 years")
        doc.save(f.name)
        yield f.name
    
    # Cleanup
    if os.path.exists(f.name):
        os.remove(f.name)


@pytest.fixture(scope="function")
def empty_pdf_file():
    """Create an empty PDF file"""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(b"")
        f.flush()
        yield f.name
    
    if os.path.exists(f.name):
        os.remove(f.name)
